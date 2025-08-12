import os
import google.generativeai as genai
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from dotenv import load_dotenv
import docx
import json
import io

# --- Environment and API Configuration ---
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found. Please set it in the .env file.")

genai.configure(api_key=GEMINI_API_KEY)

# --- Flask App Initialization ---
app = Flask(__name__)
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
app.config['TEMPLATE_DIR'] = TEMPLATE_DIR
CORS(app, resources={r"/api/*": {"origins": "http://localhost:3000"}})

# --- Helper Functions ---
def parse_docx_to_json(file_stream):
    doc = docx.Document(file_stream)
    clauses = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            clauses.append({"clause_id": f"clause_{i+1:03d}", "text": para.text.strip()})
    return clauses

# --- AI Interaction Logic ---
def get_ai_suggestions(document_json, scenario_text, form_data_text):
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
    prompt = f"""
    You are a meticulous legal AI assistant. Your task is to analyze a user's scenario against a legal document,
    incorporate pre-filled data, and suggest specific modifications.

    **Instructions:**
    1.  First, review the `base_details` which contain the foundational information for the agreement.
    2.  Next, carefully review the `scenario` which describes the specific, custom changes required.
    3.  Analyze each `clause` from the `document_clauses`.
    4.  Suggest modifications ("MODIFY", "ADD", "REMOVE") based on the `scenario`.
    5.  When generating `new_text` for "MODIFY" or "ADD" actions, you **MUST** replace the placeholders with the corresponding values from the `base_details`. Do not leave placeholders like [Client Name] in your final suggested text.
    6.  Your response **MUST** be a single, valid JSON object with one key: "suggestions".
    7.  The value of "suggestions" must be a list of JSON objects, where each object represents a single suggested change and has the following structure:
        - action: "MODIFY", "ADD", or "REMOVE".
        - clause_id: The ID of the clause to be modified or removed. For "ADD", this is the ID of the clause after which the new clause should be inserted.
        - original_text: (For "MODIFY" and "REMOVE" actions) The original text of the clause.
        - new_text: (For "MODIFY" action only) The suggested new text for the clause.
        - new_clause: (For "ADD" action only) An object with "clause_id", "clause_title", and "text" for the new clause.
        - reason: A brief, clear explanation for why the change is being suggested.

    **Base Details (from user form):**
    ---
    {form_data_text}
    ---

    **Scenario:**
    ---
    {scenario_text}
    ---

    **Document Clauses (JSON):**
    ---
    {json.dumps(document_json, indent=2)}
    ---

    **Your JSON Response:**
    """
    try:
        response = model.generate_content(prompt)
        cleaned_response_text = response.text.strip().replace("```json", "").replace("```", "")
        suggestions = json.loads(cleaned_response_text)
        return suggestions.get("suggestions", [])
    except Exception as e:
        print(f"Error during AI call or JSON parsing: {e}")
        return [{"action": "ERROR", "reason": str(e)}]

# --- API Endpoints ---
@app.route('/api/templates', methods=['GET'])
def list_templates():
    try:
        files = [f for f in os.listdir(app.config['TEMPLATE_DIR']) if f.endswith('.docx')]
        return jsonify(files)
    except FileNotFoundError:
        return jsonify({"error": "Templates directory not found."}), 404

@app.route('/api/templates/<filename>', methods=['GET'])
def get_template(filename):
    try:
        return send_from_directory(app.config['TEMPLATE_DIR'], filename, as_attachment=True)
    except FileNotFoundError:
        return jsonify({"error": "File not found."}), 404

@app.route('/api/analyze', methods=['POST'])
def analyze_document():
    if 'document' not in request.files:
        return jsonify({"error": "No document file provided"}), 400

    file = request.files['document']
    scenario = request.form.get('scenario', '')
    form_data_json = request.form.get('formData', '{}')
    
    try:
        form_data = json.loads(form_data_json)
        form_data_text = "\n".join([f"- {key.replace('_', ' ').title()}: {value}" for key, value in form_data.items()])
    except json.JSONDecodeError:
        return jsonify({"error": "Invalid form data format."}), 400

    if not scenario:
        return jsonify({"error": "No scenario text provided"}), 400

    try:
        # Use io.BytesIO to allow the stream to be read multiple times
        file_stream = io.BytesIO(file.stream.read())
        parsed_doc = parse_docx_to_json(file_stream)
        ai_suggestions = get_ai_suggestions(parsed_doc, scenario, form_data_text)

        return jsonify({
            "originalDoc": parsed_doc,
            "suggestions": ai_suggestions
        })

    except Exception as e:
        print(f"An error occurred: {e}")
        return jsonify({"error": "Failed to process the document."}), 500

# --- NEWLY ADDED CODE ---
# --- NEWLY REVISED generate_document FUNCTION ---
@app.route('/api/generate-document', methods=['POST'])
def generate_document():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid request body"}), 400

    file_name = data.get('fileName')
    suggestions = data.get('suggestions', [])
    
    # We no longer need to parse form_data here since the AI will handle it
    # in its suggestions. We'll still accept it, but it's not used in this function.
    form_data = data.get('formData', {})

    if not file_name:
        return jsonify({"error": "Missing fileName"}), 400

    try:
        template_path = os.path.join(app.config['TEMPLATE_DIR'], file_name)
        if not os.path.exists(template_path):
            return jsonify({"error": "Template file not found on server"}), 404

        # Load the original document as our template to preserve all formatting
        doc = docx.Document(template_path)
        
        # We need a robust way to find paragraphs to modify.
        # Let's map original text to the paragraph object.
        original_text_map = {p.text.strip(): p for p in doc.paragraphs if p.text.strip()}

        # A list to hold new clauses
        added_clauses_to_append = []

        # 1. Apply AI's suggestions (MODIFY, REMOVE, ADD)
        for suggestion in suggestions:
            action = suggestion.get('action')
            original_text_from_suggestion = suggestion.get('original_text', '').strip()
            new_text_from_suggestion = suggestion.get('new_text', '')

            # Find the paragraph using the original text as the key
            if action in ['MODIFY', 'REMOVE'] and original_text_from_suggestion in original_text_map:
                paragraph_to_act_on = original_text_map[original_text_from_suggestion]

                if action == 'MODIFY':
                    # Preserve original paragraph style
                    paragraph_style = paragraph_to_act_on.style
                    paragraph_to_act_on.clear()
                    paragraph_to_act_on.add_run(new_text_from_suggestion)
                    paragraph_to_act_on.style = paragraph_style

                elif action == 'REMOVE':
                    # To "remove" a paragraph, we simply clear its content
                    paragraph_to_act_on.clear()
            
            elif action == 'ADD' and 'new_clause' in suggestion:
                # Store new clauses to be appended later
                added_clauses_to_append.append(suggestion['new_clause'].get('text', ''))

        # 2. Add any new clauses to the end of the document
        for new_clause_text in added_clauses_to_append:
            if new_clause_text:
                new_paragraph = doc.add_paragraph(new_clause_text)
                # Apply a consistent style to new paragraphs
                if doc.paragraphs and len(doc.paragraphs) > 1:
                    new_paragraph.style = doc.paragraphs[-2].style

        # 3. Save the modified document to an in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # 4. Send the stream back as a file download
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"Generated_{file_name}",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Error generating document: {e}")
        return jsonify({"error": "Failed to generate the document on the server."}), 500


# --- Run the App ---
if __name__ == '__main__':
    app.run(debug=True, port=5001)